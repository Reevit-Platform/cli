from __future__ import annotations

from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
OUT = ROOT / "backend" / "docs" / "Reevit_Backend_API_Documentation.docx"
OPENAPI = BACKEND / "internal" / "docs" / "openapi.yaml"
ROUTER = BACKEND / "adapters" / "http" / "router.go"
GO_MOD = BACKEND / "go.mod"
ENV_EXAMPLE = BACKEND / ".env.example"


ACCENT = "1F4E5F"
ACCENT_2 = "0F766E"
MUTED = "64748B"
LIGHT = "EAF4F4"
TABLE_HEADER = "DDEFEF"
GREY_FILL = "F8FAFC"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    cell.width = Inches(width_inches)


def set_table_fixed(table, total_width_inches: float) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total_width_inches * 1440)))


def set_table_grid(table, widths: list[float]) -> None:
    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display" if level <= 1 else "Aptos"
        run.font.color.rgb = RGBColor.from_string(ACCENT if level <= 2 else ACCENT_2)


def add_note(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    r.font.size = Pt(10)
    p.add_run(f"\n{body}")
    for run in p.runs[1:]:
        run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    label = " / ".join(headers)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT_2)
    r.font.size = Pt(8)

    for row in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        if len(row) == 2:
            lead, body = row
            rr = p.add_run(str(lead))
            rr.bold = True
            rr.font.size = Pt(8)
            p.add_run(" - ")
            br = p.add_run(str(body))
            br.font.size = Pt(8)
        elif len(row) == 4 and headers[0].lower() == "method":
            method, path, third, fourth = row
            rr = p.add_run(f"{method} {path}")
            rr.bold = True
            rr.font.size = Pt(7)
            rr.font.color.rgb = RGBColor.from_string(ACCENT)
            rest = p.add_run(f" | {headers[2]}: {third} | {headers[3]}: {fourth}")
            rest.font.size = Pt(7)
        else:
            for i, value in enumerate(row):
                if i:
                    p.add_run(" | ")
                rr = p.add_run(f"{headers[i]}: ")
                rr.bold = True
                rr.font.size = Pt(7)
                vr = p.add_run(str(value))
                vr.font.size = Pt(7)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def parse_go_mod() -> tuple[str, list[str]]:
    text = GO_MOD.read_text()
    go_match = re.search(r"^go\s+(.+)$", text, re.M)
    deps = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith(("github.com/", "go.opentelemetry.io/", "golang.org/", "aidanwoods.dev/", "gopkg.in/")):
            deps.append(re.split(r"\s+", line)[0] + " " + re.split(r"\s+", line)[1])
    important = [
        dep for dep in deps
        if any(key in dep for key in [
            "go-chi/chi", "go-chi/cors", "pgx", "redis/go-redis", "hibiken/asynq",
            "goose", "prometheus", "opentelemetry", "go-paseto", "webauthn",
            "stripe-go", "aws-sdk-go-v2", "viper", "caarlos0/env",
        ])
    ]
    return (go_match.group(1) if go_match else "unknown", important)


def parse_openapi_paths() -> list[dict[str, str]]:
    """Small YAML-enough parser for path/method/summary/operationId from openapi.yaml."""
    entries: list[dict[str, str]] = []
    path = None
    method = None
    current = None
    for raw in OPENAPI.read_text().splitlines():
        line = raw.rstrip()
        m_path = re.match(r"^  (/[^:]+):\s*$", line)
        if m_path:
            if current:
                entries.append(current)
                current = None
            path = m_path.group(1)
            method = None
            continue
        m_method = re.match(r"^    (get|post|put|patch|delete):\s*$", line)
        if m_method and path:
            if current:
                entries.append(current)
            method = m_method.group(1).upper()
            current = {"method": method, "path": path, "summary": "", "operation": "", "tag": ""}
            continue
        if current:
            m_summary = re.match(r"^      summary:\s*(.+)$", line)
            if m_summary:
                current["summary"] = m_summary.group(1).strip().strip('"')
            m_op = re.match(r"^      operationId:\s*(.+)$", line)
            if m_op:
                current["operation"] = m_op.group(1).strip()
            m_tags = re.match(r"^      tags:\s*\[(.+)\]", line)
            if m_tags:
                current["tag"] = m_tags.group(1).strip()
    if current:
        entries.append(current)
    return entries


ENDPOINT_GROUPS = {
    "Public, Status, Docs, and Webhook Intake": [
        ("GET", "/healthz", "Public", "Readiness/liveness plus optional queue and service status."),
        ("GET", "/status/history", "Public", "Status history for public status surfaces when configured."),
        ("GET", "/status/incidents", "Public", "Public incident list for the status page."),
        ("GET", "/openapi.yaml", "Public or docs basic auth", "OpenAPI YAML specification."),
        ("GET", "/openapi.json", "Public or docs basic auth", "OpenAPI JSON specification."),
        ("GET", "/docs", "Public or docs basic auth", "Redoc documentation UI."),
        ("POST", "/webhooks/{provider}", "Provider signature", "Legacy provider webhook alias."),
        ("POST", "/v1/webhooks/incoming/{provider}", "Provider signature", "Primary PSP webhook intake."),
        ("POST", "/v1/webhooks/{provider}", "Provider signature", "Legacy v1 webhook alias."),
        ("POST", "/v1/webhooks/incoming/sailrides", "Sail Rides secret", "Delivery webhook intake."),
        ("POST", "/v1/webhooks/billing/paystack", "Paystack billing secret", "Billing webhook intake."),
        ("POST", "/v1/waitlist", "Public", "Join waitlist."),
        ("GET", "/v1/pay/{code}", "Public", "Resolve public payment link by code."),
        ("GET", "/v1/payment-links/code/{code}", "Public", "Payment-link lookup alias."),
        ("POST", "/v1/pay/{code}/pay", "Public", "Pay against a public payment link."),
        ("POST", "/v1/payments/{id}/confirm-intent", "Public", "Confirm a public payment intent."),
        ("POST", "/v1/payments/hubtel/sessions/{id}", "Public", "Create Hubtel payment session for checkout."),
    ],
    "Authentication, Identity, Session, and User Preferences": [
        ("POST", "/v1/auth/signup", "Public + rate limit", "Create user and initial organization."),
        ("POST", "/v1/auth/magic-link", "Public + rate limit", "Send magic-link sign-in email."),
        ("GET", "/v1/auth/verify", "Public + rate limit", "Exchange magic-link token for a session."),
        ("GET", "/v1/auth/verify-email-change", "Public + rate limit", "Confirm pending email address change."),
        ("POST", "/v1/auth/recover", "Public + rate limit", "Request account recovery."),
        ("GET", "/v1/auth/verify-recovery", "Public + rate limit", "Complete account recovery token flow."),
        ("GET", "/v1/auth/oauth/login", "Public", "Start Google/GitHub OAuth login."),
        ("GET", "/v1/auth/oauth/callback", "Public", "OAuth callback."),
        ("POST", "/v1/auth/oauth/link", "Session + CSRF", "Link OAuth identity to current user."),
        ("DELETE", "/v1/auth/oauth/unlink", "Session + CSRF", "Unlink OAuth identity."),
        ("GET", "/v1/auth/oauth/accounts", "Session + CSRF", "List linked OAuth accounts."),
        ("POST", "/v1/auth/refresh", "Session + CSRF", "Refresh session lifetime."),
        ("POST", "/v1/auth/logout", "Optional session + CSRF", "Invalidate session and clear cookie."),
        ("GET", "/v1/auth/me", "Session + CSRF", "Get current user profile."),
        ("PATCH", "/v1/auth/me", "Session + CSRF", "Update current user profile."),
        ("DELETE", "/v1/auth/me", "Session + CSRF", "Delete current user account."),
        ("POST", "/v1/auth/change-email", "Session + CSRF", "Start email change flow."),
        ("GET", "/v1/auth/sessions", "Session + CSRF", "List active sessions."),
        ("GET", "/v1/auth/sessions/suspicious", "Session + CSRF", "Inspect suspicious activity."),
        ("DELETE", "/v1/auth/sessions/{id}", "Session + CSRF", "Revoke one session."),
        ("POST", "/v1/auth/sessions/revoke-others", "Session + CSRF", "Revoke other sessions."),
        ("POST", "/v1/auth/2fa/setup", "Session + CSRF", "Create TOTP secret and QR setup payload."),
        ("POST", "/v1/auth/2fa/enable", "Session + CSRF", "Enable TOTP after code verification."),
        ("POST", "/v1/auth/2fa/disable", "Session + CSRF", "Disable TOTP."),
        ("GET", "/v1/auth/2fa/status", "Session + CSRF", "Read TOTP status."),
        ("POST", "/v1/auth/2fa/backup-codes", "Session + CSRF", "Regenerate backup codes."),
        ("GET", "/v1/auth/preferences", "Session + CSRF", "Get notification/user preferences."),
        ("PATCH", "/v1/auth/preferences", "Session + CSRF", "Update preferences."),
        ("GET", "/v1/auth/activities", "Session + CSRF", "List account activity log."),
        ("POST", "/v1/auth/avatar/presigned", "Session + CSRF", "Create presigned avatar upload URL."),
        ("PATCH", "/v1/auth/me/avatar", "Session + CSRF", "Attach avatar URL to user."),
        ("DELETE", "/v1/auth/me/avatar", "Session + CSRF", "Remove avatar."),
        ("POST", "/v1/auth/webauthn/register/begin", "Session + CSRF", "Begin passkey registration."),
        ("POST", "/v1/auth/webauthn/register/finish", "Session + CSRF", "Finish passkey registration."),
        ("POST", "/v1/auth/webauthn/authenticate/begin", "Session + CSRF", "Begin passkey authentication."),
        ("POST", "/v1/auth/webauthn/authenticate/finish", "Session + CSRF", "Finish passkey authentication."),
        ("GET", "/v1/auth/webauthn/credentials", "Session + CSRF", "List passkeys."),
        ("DELETE", "/v1/auth/webauthn/credentials", "Session + CSRF", "Delete passkey."),
    ],
    "Onboarding, KYC, Platform, and Admin Operations": [
        ("GET", "/v1/features/enabled", "Session + CSRF", "List dashboard features enabled for the user/org."),
        ("POST", "/v1/onboarding/start", "Session + CSRF", "Start onboarding."),
        ("GET", "/v1/onboarding/form", "Session + CSRF", "Fetch onboarding form schema."),
        ("POST", "/v1/onboarding/answers", "Session + CSRF", "Save onboarding answers."),
        ("POST", "/v1/onboarding/upgrade-to-business", "Session + CSRF", "Upgrade individual onboarding to business."),
        ("DELETE", "/v1/onboarding/draft", "Session + CSRF", "Delete onboarding draft."),
        ("POST", "/v1/kyc/submit", "Session + org + CSRF", "Submit organization KYC."),
        ("GET", "/v1/kyc/status", "Session + org + CSRF", "Get organization KYC status."),
        ("GET", "/v1/platform/me", "Session + CSRF", "Check current user's platform admin role."),
        ("GET", "/v1/platform/kyc/submissions", "Platform admin", "List KYC submissions."),
        ("POST", "/v1/platform/kyc/approve", "Platform admin", "Approve KYC."),
        ("POST", "/v1/platform/kyc/reject", "Platform admin", "Reject KYC."),
        ("GET", "/v1/platform/kyc/{org_id}/notes", "Platform admin", "List internal KYC notes."),
        ("POST", "/v1/platform/kyc/{org_id}/notes", "Platform admin", "Add internal KYC note."),
        ("GET", "/v1/platform/organizations", "Platform admin", "List organizations."),
        ("GET", "/v1/platform/organizations/{id}", "Platform admin", "Get organization."),
        ("PATCH", "/v1/platform/organizations/{id}", "Platform admin", "Update organization."),
        ("DELETE", "/v1/platform/organizations/{id}", "Platform admin", "Delete organization."),
        ("GET", "/v1/platform/audit-logs", "Platform admin", "List platform audit events."),
        ("GET", "/v1/platform/users", "Platform admin", "List users."),
        ("GET", "/v1/platform/users/{id}", "Platform admin", "Get user."),
        ("PATCH", "/v1/platform/users/{id}", "Platform admin", "Update user."),
        ("DELETE", "/v1/platform/users/{id}", "Platform admin", "Delete user."),
        ("POST", "/v1/platform/users/{id}/suspend", "Platform admin", "Suspend user."),
        ("POST", "/v1/platform/users/{id}/unsuspend", "Platform admin", "Unsuspend user."),
        ("GET", "/v1/platform/users/{id}/suspension", "Platform admin", "Read suspension status."),
        ("DELETE", "/v1/platform/users/{id}/organizations/{org_id}", "Platform admin", "Remove user from org."),
        ("GET", "/v1/platform/admins", "Platform owner", "List platform admins."),
        ("POST", "/v1/platform/admins", "Platform owner", "Add platform admin."),
        ("DELETE", "/v1/platform/admins", "Platform owner", "Remove platform admin."),
        ("GET", "/v1/platform/plans", "Platform admin", "List billing plans."),
        ("POST", "/v1/platform/plans", "Platform admin", "Create billing plan."),
        ("GET", "/v1/platform/plans/{id}", "Platform admin", "Get billing plan."),
        ("PUT", "/v1/platform/plans/{id}", "Platform admin", "Update billing plan."),
        ("POST", "/v1/platform/plans/{id}/archive", "Platform admin", "Archive plan."),
        ("POST", "/v1/platform/plans/{id}/restore", "Platform admin", "Restore plan."),
        ("DELETE", "/v1/platform/plans/{id}", "Platform admin", "Delete plan."),
        ("GET", "/v1/platform/subscriptions", "Platform admin", "List org subscriptions."),
        ("GET", "/v1/platform/coupons", "Platform admin", "List coupons."),
        ("POST", "/v1/platform/coupons", "Platform admin", "Create coupon."),
        ("GET", "/v1/platform/coupons/analytics", "Platform admin", "Coupon analytics."),
        ("GET", "/v1/platform/settings", "Platform owner", "Read platform settings."),
        ("PUT", "/v1/platform/settings/{key}", "Platform owner", "Update platform setting."),
        ("GET", "/v1/platform/features", "Platform admin", "List dashboard feature flags."),
        ("PATCH", "/v1/platform/features/{key}", "Platform admin", "Update feature flag."),
        ("PATCH", "/v1/platform/features/bulk", "Platform admin", "Bulk update feature flags."),
        ("GET", "/v1/platform/incidents", "Platform admin", "List incidents."),
        ("POST", "/v1/platform/incidents", "Platform admin", "Create incident."),
        ("GET", "/v1/platform/waitlist", "Platform admin", "List waitlist entries."),
        ("GET", "/v1/platform/failure-events", "Platform admin", "List platform failure events."),
        ("GET", "/v1/platform/logs", "Platform admin", "List platform logs."),
        ("GET", "/v1/platform/traces", "Platform admin", "List traces."),
        ("GET", "/v1/admin/users", "Session + admin", "List organization members."),
        ("POST", "/v1/admin/users/invite", "Session + admin", "Invite user to organization."),
        ("POST", "/v1/admin/users/bulk/roles", "Session + admin", "Bulk role update."),
        ("GET", "/v1/admin/organizations", "Session + owner", "Get current organization."),
        ("PATCH", "/v1/admin/organizations", "Session + owner", "Update current organization."),
        ("POST", "/v1/admin/organizations/logo/presigned", "Session + owner", "Create presigned org logo upload URL."),
        ("GET", "/v1/admin/audit-logs", "Session + owner", "Organization audit trail."),
    ],
    "Merchant API, Payments, Billing, Orchestration, and Operations": [
        ("GET", "/v1/events/stream", "Session or API key + org", "Server-sent event stream."),
        ("GET", "/v1/billing/subscription", "Session or API key + org + CSRF", "Current billing subscription."),
        ("GET", "/v1/billing/usage", "Session or API key + org + CSRF", "Plan usage/meters."),
        ("GET", "/v1/billing/plans", "Session or API key + org + CSRF", "Available plans."),
        ("POST", "/v1/billing/subscription", "Session or API key + org + CSRF", "Update subscription."),
        ("POST", "/v1/billing/checkout", "Session or API key + org + CSRF", "Create Paystack billing checkout."),
        ("POST", "/v1/billing/coupons/validate", "Session or API key + org + CSRF", "Validate coupon."),
        ("POST", "/v1/billing/coupons/apply", "Session or API key + org + CSRF", "Apply coupon."),
        ("GET", "/v1/api-keys", "api_keys:read", "List API keys."),
        ("POST", "/v1/api-keys", "api_keys:write + idempotency", "Create API key."),
        ("DELETE", "/v1/api-keys/{id}", "api_keys:write", "Revoke API key."),
        ("GET", "/v1/connections", "connections:read", "List PSP connections."),
        ("POST", "/v1/connections", "connections:write + idempotency", "Upsert PSP connection."),
        ("GET", "/v1/connections/{id}", "connections:read", "Get PSP connection."),
        ("DELETE", "/v1/connections/{id}", "connections:write", "Delete PSP connection."),
        ("PATCH", "/v1/connections/{id}/labels", "connections:write", "Update connection labels."),
        ("PATCH", "/v1/connections/{id}/status", "connections:write", "Update connection status."),
        ("POST", "/v1/connections/{id}/validate", "connections:write + idempotency", "Validate credentials."),
        ("POST", "/v1/connections/test", "connections:write + idempotency", "Test credentials before saving."),
        ("GET", "/v1/checkout/settings", "payments:read", "Read checkout settings."),
        ("PUT", "/v1/checkout/settings", "payments:write + idempotency", "Upsert checkout settings."),
        ("GET", "/v1/routing-rules", "routing_rules:read + smart_routing", "List routing rules."),
        ("POST", "/v1/routing-rules", "routing_rules:write + smart_routing", "Create routing rule."),
        ("PATCH", "/v1/routing-rules/{id}", "routing_rules:write + smart_routing", "Update routing rule."),
        ("GET", "/v1/routing/ab-tests", "routing_rules:read + smart_routing", "List routing A/B tests."),
        ("POST", "/v1/routing/ab-tests", "routing_rules:write + smart_routing", "Create routing A/B test."),
        ("GET", "/v1/routing/decisions", "routing_rules:read + smart_routing", "List recorded routing decisions."),
        ("GET", "/v1/customers", "customers:read", "List customer profiles."),
        ("POST", "/v1/customers", "customers:write + idempotency", "Create customer profile."),
        ("GET", "/v1/customers/top", "customers:read", "Top customers."),
        ("GET", "/v1/customers/lookup", "customers:read", "Lookup customer by external ID."),
        ("GET", "/v1/payment-links", "payment_links:read", "List payment links."),
        ("POST", "/v1/payment-links", "payment_links:write + idempotency", "Create payment link."),
        ("GET", "/v1/payment-links/{id}/stats", "payment_links:read", "Payment link stats."),
        ("GET", "/v1/dunning/campaigns", "subscriptions:read + basic_failover", "List dunning campaigns."),
        ("POST", "/v1/dunning/campaigns", "subscriptions:write + basic_failover", "Create dunning campaign."),
        ("GET", "/v1/payments", "payments:read", "List payments."),
        ("POST", "/v1/payments/intents", "payments:write + idempotency", "Create payment intent."),
        ("PATCH", "/v1/payments/intents/{id}", "payments:write + idempotency", "Update payment intent."),
        ("POST", "/v1/payments/{id}/confirm", "payments:write + idempotency", "Confirm payment."),
        ("POST", "/v1/payments/{id}/capture", "payments:write + idempotency", "Capture authorized payment."),
        ("POST", "/v1/payments/{id}/refund", "payments:write + idempotency", "Refund payment."),
        ("POST", "/v1/payments/{id}/retry", "payments:write + idempotency", "Retry payment."),
        ("POST", "/v1/payments/{id}/cancel", "payments:write + idempotency", "Cancel payment intent."),
        ("GET", "/v1/payments/stats", "payments:read", "Payment analytics summary."),
        ("GET", "/v1/payments/stats/series", "payments:read", "Payment time series."),
        ("GET", "/v1/payments/stats/errors", "payments:read", "Error analytics."),
        ("GET", "/v1/payments/stats/routing", "payments:read", "Routing analytics."),
        ("GET", "/v1/payments/refunds", "payments:read", "List refunds."),
        ("GET", "/v1/payments/disputes", "payments:read", "List disputes."),
        ("POST", "/v1/payments/disputes/{id}/resolve", "payments:write + idempotency", "Resolve dispute."),
        ("POST", "/v1/payments/scheduled", "payments:write + idempotency", "Create scheduled payment."),
        ("GET", "/v1/payment-methods", "payments:read", "List saved payment methods."),
        ("POST", "/v1/payment-methods", "payments:write + idempotency", "Create saved payment method."),
        ("GET", "/v1/policies/fraud", "fraud:read + fraud feature", "Read fraud policy."),
        ("POST", "/v1/policies/fraud", "fraud:write + fraud feature", "Upsert fraud policy."),
        ("GET", "/v1/policies/retry", "subscriptions:read + failover", "Read retry policy."),
        ("POST", "/v1/policies/retry", "subscriptions:write + failover", "Upsert retry policy."),
        ("GET", "/v1/webhooks/events", "webhooks:read + webhooks feature", "List inbound webhook events."),
        ("POST", "/v1/webhooks/events/{event_id}/replay", "webhooks:write + webhooks feature", "Replay inbound webhook event."),
        ("GET", "/v1/webhooks/config", "webhooks:read + webhooks feature", "Read outbound webhook config."),
        ("POST", "/v1/webhooks/config", "webhooks:write + webhooks feature", "Upsert outbound webhook config."),
        ("GET", "/v1/webhooks/dashboard", "webhooks:read + webhooks feature", "Webhook delivery dashboard."),
        ("GET", "/v1/subscriptions", "subscriptions:read", "List subscriptions."),
        ("POST", "/v1/subscriptions", "subscriptions:write + idempotency", "Create subscription."),
        ("PATCH", "/v1/subscriptions/{id}", "subscriptions:write + idempotency", "Update subscription."),
        ("POST", "/v1/subscriptions/{id}/cancel", "subscriptions:write + idempotency", "Cancel subscription."),
        ("GET", "/v1/invoices", "invoices:read", "List invoices."),
        ("PATCH", "/v1/invoices/{id}", "invoices:write + idempotency", "Update invoice metadata."),
        ("POST", "/v1/invoices/{id}/retry", "invoices:write + idempotency", "Retry invoice payment."),
        ("GET", "/v1/workflows/integrations", "workflows:read + workflows feature", "List workflow integrations."),
        ("POST", "/v1/workflows/rules", "workflows:write + workflows feature", "Create workflow rule."),
        ("POST", "/v1/workflows/rules/{id}/publish", "workflows:write + workflows feature", "Publish workflow version."),
        ("GET", "/v1/workflows/executions", "workflows:read + workflows feature", "List workflow executions."),
        ("GET", "/v1/exports/payments", "payments:read", "Export payments."),
        ("GET", "/v1/health/connections", "connections:read", "Connection health dashboard."),
        ("GET", "/v1/analytics/api-usage", "api_keys:read", "API usage analytics."),
        ("GET", "/v1/failure-events", "payments:read", "Failure event dashboard."),
        ("GET", "/v1/notifications", "Session or API key + org", "List notifications."),
    ],
}


def build_doc() -> None:
    go_version, deps = parse_go_mod()
    openapi_entries = parse_openapi_paths()
    router_text = ROUTER.read_text()
    migration_count = len(list((BACKEND / "db" / "migrations").glob("*.sql")))
    query_count = len(list((BACKEND / "db" / "queries").glob("*.sql")))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(26)
    styles["Title"].font.bold = True
    for style in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style].font.name = "Aptos Display" if style == "Heading 1" else "Aptos"
        styles[style].font.color.rgb = RGBColor.from_string(ACCENT)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Reevit Backend API Documentation")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Architecture, endpoint catalog, auth model, data model, and operations guide")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    meta_rows = [
        ["Project", "primeflow / backend (Reevit Platform)"],
        ["Generated", date.today().isoformat()],
        ["Primary source", "backend/adapters/http/router.go"],
        ["Secondary source", "backend/internal/docs/openapi.yaml and backend/docs/internal/*.md"],
        ["Language", f"Go {go_version}"],
        ["Endpoint count from OpenAPI", str(len(openapi_entries))],
        ["Database artifacts", f"{migration_count} migrations, {query_count} SQLC query files"],
    ]
    add_table(doc, ["Field", "Value"], meta_rows, [1.8, 5.7])

    add_note(
        doc,
        "Documentation stance",
        "This manual is generated from the backend implementation and existing project documentation. "
        "Where router.go contains newer endpoints than OpenAPI, the implementation route is treated as authoritative.",
    )

    doc.add_page_break()
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "The backend is a Go-based, clean/hexagonal payment orchestration API for Reevit. It consolidates PSP "
        "connections, payment intents, subscriptions, invoices, retry policy, fraud policy, routing rules, customer "
        "profiles, webhooks, workflows, platform administration, and operational observability behind a single API."
    )
    bullets(doc, [
        "Core runtime: Go HTTP service using chi, PostgreSQL, Redis, Asynq, SQLC, Goose, Prometheus, and OpenTelemetry.",
        "Business model: BYOK PSP orchestration. Merchants connect their own Paystack, Hubtel, Flutterwave, Monnify, M-Pesa, and Stripe accounts.",
        "Security model: API keys with scopes for server-to-server calls, session cookies or Bearer sessions for dashboard users, CSRF protection for session mutations, PASETO tokens, optional TOTP and WebAuthn.",
        "Tenant isolation: most protected endpoints require X-Org-Id and are guarded by organization membership, role, scope, and plan-feature middleware.",
        "Operational posture: health endpoints, queue metrics, Prometheus, OpenTelemetry traces, failure-event dashboards, platform logs, status incidents, and webhook delivery dashboards.",
    ])

    add_heading(doc, "2. Backend Stack", 1)
    add_table(doc, ["Layer", "Implementation"], [
        ["HTTP", "net/http plus go-chi/chi router and go-chi/cors"],
        ["Authentication", "API keys, dashboard sessions, PASETO, OAuth, TOTP, WebAuthn"],
        ["Data", "PostgreSQL via pgx pool, SQLC generated query layer, Goose migrations"],
        ["Cache/queue", "Redis for idempotency/rate/usage support; Asynq for jobs"],
        ["Observability", "Prometheus metrics, expvar, OpenTelemetry traces"],
        ["Secrets", "AES vault with optional AWS KMS support"],
        ["PSP adapters", "Paystack, Hubtel, Flutterwave, Monnify, M-Pesa, Stripe, stub; additional provider stubs exist for Interswitch and OPay"],
        ["Docs", "Embedded OpenAPI YAML/JSON and Redoc UI"],
    ], [1.7, 5.9])
    add_table(doc, ["Important library", "Version"], [dep.split(" ", 1) for dep in deps], [4.7, 2.4])

    add_heading(doc, "3. Architecture", 1)
    doc.add_paragraph(
        "The backend follows a clean/hexagonal layout. The HTTP adapter converts request/response DTOs, use cases hold "
        "business flows, ports define contracts, repositories and PSP adapters implement those contracts, and cmd/api wires "
        "runtime dependencies. The worker command processes scheduled and asynchronous work."
    )
    add_table(doc, ["Directory", "Responsibility"], [
        ["/cmd/api", "HTTP server bootstrap, dependency wiring, metrics endpoints, graceful shutdown."],
        ["/cmd/worker", "Asynq worker and scheduler for renewals, dunning, webhook fanout, scheduled payments, and side effects."],
        ["/adapters/http", "Chi router, handlers, response helpers, auth/rate/scope/mode/plan/idempotency middleware."],
        ["/adapters/repo", "PostgreSQL repositories implementing ports using SQLC and pgx."],
        ["/adapters/psp", "Provider-specific payment adapters and credential testers."],
        ["/adapters/webhook", "Provider webhook parsers normalizing inbound events."],
        ["/internal/usecase", "Payments, connections, billing, subscriptions, routing, exports, webhooks, workflows, health, incidents, notifications."],
        ["/internal/services", "Auth, sessions, TOTP, OAuth, WebAuthn, credentials, storage, usage buffer, webhook retry helpers."],
        ["/internal/infra", "Configuration, database, idempotency stores, HTTP client, vault, telemetry, logging."],
        ["/db/migrations and /db/queries", "Goose schema evolution and SQLC query definitions."],
    ], [2.0, 5.6])

    add_heading(doc, "4. Request Conventions", 1)
    add_table(doc, ["Convention", "Details"], [
        ["Base paths", "Runtime exposes root endpoints plus /v1 business APIs. OpenAPI lists localhost:8081 and api.reevit.io; .env.example defaults SERVER_ADDR=:8080."],
        ["Content type", "JSON request and response bodies unless an export endpoint returns a file format."],
        ["Tenant header", "X-Org-Id is required for protected tenant data access."],
        ["API key header", "X-Reevit-Key: pfk_<org_id>.<secret>; scopes are enforced per endpoint group."],
        ["Session", "Cookie reevit_session or Authorization: Bearer <session_token>."],
        ["Idempotency", "Idempotency-Key is supported on mutating endpoints that attach the idempotency middleware; default TTL is 24 hours."],
        ["Money", "Amounts are int64 minor units; do not send floating-point amounts."],
        ["Mode", "Several data sets are mode-aware: live/test isolation is enforced in repository and middleware paths."],
        ["Limits", "Global body limit is 10MB; webhook payload limit is stricter, configured by ContentLimitConfig."],
    ], [2.0, 5.6])

    add_heading(doc, "5. Security and Middleware", 1)
    bullets(doc, [
        "Global security headers include nosniff, frame deny, XSS protection, referrer policy, and a restrictive CSP.",
        "CORS is allowlist-driven. Production panics when FRONTEND_ALLOWED_ORIGINS is not explicitly configured.",
        "Session-authenticated mutations use CSRF middleware. API-key calls are scope-gated.",
        "Admin routes can require IP whitelist, rate limiting, 2FA, and audit logging.",
        "Plan feature gates protect smart routing, workflows, webhooks, fraud rules, failover, audit logs, and live mode access.",
        "Rate limiting exists for public auth endpoints, authenticated API traffic, and admin traffic.",
        "Webhook routes use provider signature parsing plus smaller payload limits.",
    ])
    add_table(doc, ["Scope", "Primary endpoint families"], [
        ["connections:read/write", "Connections, connection health, connection failures, credential validation/test."],
        ["payments:read/write", "Payment intents, confirmations, capture, refunds, disputes, fees, analytics, scheduled payments, payment methods."],
        ["fraud:read/write", "Fraud policy and audit."],
        ["webhooks:read/write", "Webhook events, outbound config, replay, dashboard deliveries."],
        ["subscriptions:read/write", "Subscriptions, invoices, retry policy, dunning, related exports."],
        ["api_keys:read/write", "API-key management and usage analytics."],
        ["workflows:read/write", "Workflow integrations, templates, rules, executions, fraud rules, versioning."],
        ["routing_rules:read/write", "Routing rules, A/B tests, routing decision audit."],
        ["customers:read/write", "Customer profiles and payment history."],
        ["payment_links:read/write", "Payment-link management, stats, linked payments."],
    ], [2.1, 5.5])

    add_heading(doc, "6. Endpoint Catalog", 1)
    for group, rows in ENDPOINT_GROUPS.items():
        add_heading(doc, group, 2)
        add_table(doc, ["Method", "Path", "Auth/scope", "Purpose"], [list(r) for r in rows], [0.7, 2.6, 1.8, 3.0])

    add_heading(doc, "7. OpenAPI Operation Index", 1)
    doc.add_paragraph(
        "The following index is parsed from backend/internal/docs/openapi.yaml. It is useful for SDK generation and "
        "client-facing operation names. Some implementation-only routes appear in the catalog above before they appear in OpenAPI."
    )
    api_rows = [
        [e["method"], e["path"], e.get("operation", ""), e.get("summary", "")]
        for e in openapi_entries
        if e["path"].startswith("/v1/") or e["path"] in ["/healthz", "/metrics"]
    ]
    add_table(doc, ["Method", "Path", "Operation ID", "Summary"], api_rows, [0.6, 2.7, 1.8, 2.7])

    add_heading(doc, "8. Domain and Data Model", 1)
    add_table(doc, ["Domain", "Key entities and tables"], [
        ["Identity", "users, sessions, memberships, organizations, oauth_accounts, webauthn_credentials, preferences, activity logs."],
        ["API access", "api_keys with encrypted/hashed secret material, scopes, revocation, usage analytics."],
        ["PSP connections", "connections, connection audit logs, health buckets, validation details, fee structures, labels, routing hints."],
        ["Payments", "payments, refunds, payment_status_history, payment_disputes, saved_payment_methods, scheduled_payments, side-effect outbox."],
        ["Subscriptions", "subscriptions, invoices, retry_policies, dunning campaigns and attempts."],
        ["Routing", "routing_rules, routing_ab_tests, routing_decisions; telemetry-aware provider selection."],
        ["Webhooks", "webhook_events, org_webhook_configs, outbound_webhooks/deliveries with replay and retry."],
        ["Billing/platform", "plans, org_subscriptions, invoices, coupons, platform_admins, platform_audit_logs, platform settings."],
        ["Operations", "failure_events, notifications, log_entries, traces, incidents, waitlist entries, KYC notes."],
        ["Workflows", "workflow integrations, templates, rules, versions, history, execution/audit records, custom fraud rules."],
    ], [1.7, 5.9])
    add_note(
        doc,
        "Composite-key caution",
        "The connections table is documented as org-scoped with composite access patterns: org-scoped API operations should use org_id plus connection_id; system-level workers and webhooks may use connection ID lookups only where explicitly intended.",
    )

    add_heading(doc, "9. Payment and Routing Behavior", 1)
    numbered(doc, [
        "Merchant creates PSP connections using BYOK credentials. Secrets are encrypted before persistence.",
        "Payment intent creation validates money, method, currency, country, optional customer, metadata, and routing preferences.",
        "The routing layer evaluates explicit connection preference, country/method rules, connection health, fee hints, feature gates, and provider capabilities.",
        "The selected PSP adapter creates or confirms the payment with provider-specific payloads.",
        "Provider webhooks are parsed into a unified event model and update payments, status history, metrics, outbound webhooks, and failure events.",
        "Retry, capture, cancel, refund, scheduled payment, and dunning flows reuse the same tenant, scope, idempotency, and side-effect rules.",
    ])
    add_table(doc, ["Provider", "Role in backend"], [
        ["Paystack", "Cards, bank/USSD, billing checkout and billing webhooks, credential testing."],
        ["Hubtel", "Ghana MoMo/card flows and Hubtel checkout sessions."],
        ["Flutterwave", "Fallback card/MoMo/bank flows and webhook parsing."],
        ["Monnify", "Bank-transfer-oriented provider support."],
        ["M-Pesa", "Mobile-money provider support."],
        ["Stripe", "Card rails including wallet-backed PaymentIntent flows; webhook metadata links events to org/connection/payment."],
        ["Stub", "Development/test provider adapter."],
    ], [1.3, 6.2])

    add_heading(doc, "10. Workers, Events, and Observability", 1)
    bullets(doc, [
        "cmd/worker uses Asynq and scheduler code to process renewals, dunning, webhook delivery, scheduled payments, and other deferred side effects.",
        "internal/eventstream provides an in-process event broker for Server-Sent Events at /v1/events/stream when enabled.",
        "Prometheus metrics are exposed through metrics wiring in cmd/api and provider/routing/billing/webhook/auth metric recorders.",
        "OpenTelemetry traces are initialized from environment-backed telemetry config and attached through tracing middleware.",
        "Failure events, platform logs, and traces have both merchant-facing and platform-admin-facing dashboards.",
        "Health surfaces include /healthz, public status history/incidents, connection health dashboards, and webhook delivery dashboards.",
    ])

    add_heading(doc, "11. Environment and Runtime Configuration", 1)
    env_keys = []
    for line in ENV_EXAMPLE.read_text().splitlines():
        if line and not line.startswith("#") and "=" in line:
            key = line.split("=", 1)[0]
            if key:
                env_keys.append(key)
    add_table(doc, ["Configuration area", "Variables"], [
        ["Application", ", ".join(k for k in env_keys if k in ["APP_ENV", "SERVER_ADDR", "DASHBOARD_BASE_URL", "PUBLIC_API_BASE_URL"])],
        ["Database", ", ".join(k for k in env_keys if k.startswith("POSTGRES") or k == "DATABASE_URL")],
        ["Security", ", ".join(k for k in env_keys if k in ["VAULT_KEY", "PASETO_KEY", "SSRF_ALLOWED_DOMAINS", "COOKIE_DOMAIN"])],
        ["Redis/idempotency", ", ".join(k for k in env_keys if k in ["REDIS_ADDR", "IDEMPOTENCY_TTL"])],
        ["Email/auth", ", ".join(k for k in env_keys if k.startswith(("RESEND", "NOTIFIER_EMAIL", "OAUTH_", "WEBAUTHN_")))],
        ["Admin/rate", ", ".join(k for k in env_keys if k.startswith(("ADMIN_", "REQUIRE_2FA", "PUBLIC_AUTH")))],
        ["Providers", ", ".join(k for k in env_keys if k.startswith(("PAYSTACK", "HUBTEL", "FLUTTERWAVE", "MONNIFY", "MPESA")))],
        ["Docs/status", ", ".join(k for k in env_keys if k.startswith(("DOCS_", "STATUS_")))],
    ], [1.8, 5.8])

    add_heading(doc, "12. Implementation Notes and Gaps to Watch", 1)
    bullets(doc, [
        "Router.go is broader than the existing internal markdown route reference, especially around platform operations, coupons, checkout settings, payment links, dunning, notifications, failure events, and workflow versioning.",
        "OpenAPI should be regenerated or expanded to include every implementation route if SDKs and external docs are expected to be complete.",
        "Some route availability is conditional on dependency wiring. For example, platform logs, traces, waitlist, incidents, KYC notes, checkout settings, event stream, and scheduled payments only mount when corresponding services/stores are configured.",
        "Session endpoints often include CSRF protection even on GET in the router chain; dashboard clients should preserve the CSRF/session conventions used by the frontend.",
        "Production deployment must set explicit CORS origins, secure VAULT_KEY/PASETO_KEY, provider secrets, cookie domain, and webhook signing secrets.",
        "Tenant safety depends on X-Org-Id, membership checks, mode isolation, and org-scoped repository methods. Avoid introducing system-level lookups into user-facing paths unless deliberately reviewed.",
    ])

    # Landscape appendix for wide route verification snippets.
    doc.add_section(WD_ORIENT.LANDSCAPE)
    add_heading(doc, "Appendix A. Router Route Count Evidence", 1)
    doc.add_paragraph(
        "The generator also inspected router.go for mounted methods and middleware. This appendix records high-level counts "
        "from the implementation source so reviewers can reconcile this document with the code."
    )
    method_counts = []
    for method in ["Get", "Post", "Put", "Patch", "Delete"]:
        method_counts.append([method.upper(), str(len(re.findall(rf"\.{method}\(", router_text)))])
    add_table(doc, ["Chi method", "Occurrences in router.go"], method_counts, [2.0, 2.0])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
