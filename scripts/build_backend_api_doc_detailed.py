from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path
import json
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
OPENAPI = BACKEND / "internal" / "docs" / "openapi.yaml"
ROUTER = BACKEND / "adapters" / "http" / "router.go"
GO_MOD = BACKEND / "go.mod"
ENV_EXAMPLE = BACKEND / ".env.example"
OUT = BACKEND / "docs" / "Reevit_Backend_API_Documentation.docx"

ACCENT = "1F4E5F"
GREEN = "0F766E"
MUTED = "64748B"
RED = "991B1B"


def line_indent(line: str) -> int:
    return len(line) - len(line.lstrip(" "))


def clean_scalar(value: str) -> str:
    value = value.strip()
    if value.startswith('"') and value.endswith('"'):
        return value[1:-1]
    if value.startswith("'") and value.endswith("'"):
        return value[1:-1]
    return value


def parse_go_mod() -> tuple[str, list[tuple[str, str]]]:
    text = GO_MOD.read_text()
    go_match = re.search(r"^go\s+(.+)$", text, re.M)
    wanted = [
        "go-paseto", "go-chi/chi", "go-chi/cors", "pgx/v5", "go-redis",
        "asynq", "goose", "prometheus", "opentelemetry", "webauthn",
        "stripe-go", "aws-sdk-go-v2", "viper", "caarlos0/env",
        "oauth2", "crypto",
    ]
    deps: list[tuple[str, str]] = []
    for line in text.splitlines():
        line = line.strip()
        parts = line.split()
        if len(parts) < 2:
            continue
        if any(w in parts[0] for w in wanted):
            deps.append((parts[0], parts[1]))
    return go_match.group(1) if go_match else "unknown", deps


def parse_schema_ref(line: str) -> str | None:
    m = re.search(r"#/components/schemas/([A-Za-z0-9_]+)", line)
    return m.group(1) if m else None


def parse_response_ref(line: str) -> str | None:
    m = re.search(r"#/components/responses/([A-Za-z0-9_]+)", line)
    return m.group(1) if m else None


def parse_param_ref(line: str) -> str | None:
    m = re.search(r"#/components/parameters/([A-Za-z0-9_]+)", line)
    return m.group(1) if m else None


def parse_paths() -> list[dict]:
    lines = OPENAPI.read_text().splitlines()
    entries: list[dict] = []
    current_path = None
    i = 0
    while i < len(lines):
        line = lines[i]
        m_path = re.match(r"^  (/[^:]+):\s*$", line)
        if m_path:
            current_path = m_path.group(1)
            i += 1
            continue
        m_method = re.match(r"^    (get|post|put|patch|delete):\s*$", line)
        if current_path and m_method:
            method = m_method.group(1).upper()
            start = i
            i += 1
            while i < len(lines):
                if re.match(r"^  (/[^:]+):\s*$", lines[i]) or re.match(r"^    (get|post|put|patch|delete):\s*$", lines[i]):
                    break
                if lines[i].startswith("components:"):
                    break
                i += 1
            block = lines[start:i]
            entries.append(parse_operation(current_path, method, block))
            continue
        if line.startswith("components:"):
            break
        i += 1
    return entries


def parse_operation(path: str, method: str, block: list[str]) -> dict:
    op = {
        "path": path,
        "method": method,
        "summary": "",
        "description": "",
        "operationId": "",
        "tags": [],
        "parameters": [],
        "request_schema": "",
        "request_required": False,
        "responses": [],
    }
    i = 0
    while i < len(block):
        line = block[i]
        stripped = line.strip()
        if stripped.startswith("summary:"):
            op["summary"] = clean_scalar(stripped.split(":", 1)[1])
        elif stripped.startswith("operationId:"):
            op["operationId"] = clean_scalar(stripped.split(":", 1)[1])
        elif stripped.startswith("tags:"):
            inside = stripped.split(":", 1)[1].strip()
            op["tags"] = [x.strip() for x in inside.strip("[]").split(",") if x.strip()]
        elif stripped.startswith("description: |"):
            desc_lines = []
            base_indent = line_indent(line)
            i += 1
            while i < len(block) and line_indent(block[i]) > base_indent:
                desc_lines.append(block[i].strip())
                i += 1
            op["description"] = "\n".join(desc_lines).strip()
            continue
        elif stripped.startswith("- $ref:") and "parameters" in "\n".join(block[max(0, i - 3):i + 1]):
            ref = parse_param_ref(stripped)
            if ref:
                op["parameters"].append({"name": ref, "in": "header", "required": "varies", "type": "ref", "description": f"Reusable parameter {ref}"})
        elif re.match(r"-\s+in:\s+", stripped):
            param = {"name": "", "in": clean_scalar(stripped.split(":", 1)[1]), "required": "false", "type": "", "description": ""}
            base_indent = line_indent(line)
            i += 1
            while i < len(block) and line_indent(block[i]) > base_indent:
                s = block[i].strip()
                if s.startswith("name:"):
                    param["name"] = clean_scalar(s.split(":", 1)[1])
                elif s.startswith("required:"):
                    param["required"] = clean_scalar(s.split(":", 1)[1])
                elif s.startswith("description:"):
                    param["description"] = clean_scalar(s.split(":", 1)[1])
                elif s.startswith("type:"):
                    param["type"] = clean_scalar(s.split(":", 1)[1])
                elif s.startswith("default:") and param["description"]:
                    param["description"] += f" Default: {clean_scalar(s.split(':', 1)[1])}."
                elif s.startswith("enum:"):
                    param["description"] += f" Allowed: {clean_scalar(s.split(':', 1)[1])}."
                i += 1
            if param["name"]:
                op["parameters"].append(param)
            continue
        elif stripped.startswith("requestBody:"):
            body_block = collect_child_block(block, i)
            op["request_required"] = any("required: true" in x for x in body_block)
            for b in body_block:
                ref = parse_schema_ref(b)
                if ref:
                    op["request_schema"] = ref
                    break
        elif re.match(r"'?[0-9]{3}'?:", stripped):
            code = stripped.split(":", 1)[0].strip("'")
            response_block = collect_child_block(block, i)
            desc = ""
            schema = ""
            resp_ref = ""
            for b in response_block:
                bs = b.strip()
                if bs.startswith("$ref:"):
                    resp_ref = parse_response_ref(bs) or ""
                elif bs.startswith("description:") and not desc:
                    desc = clean_scalar(bs.split(":", 1)[1])
                schema = schema or (parse_schema_ref(bs) or "")
            op["responses"].append({"code": code, "description": desc, "schema": schema, "ref": resp_ref})
        i += 1
    return op


def collect_child_block(block: list[str], index: int) -> list[str]:
    base = line_indent(block[index])
    out = []
    j = index + 1
    while j < len(block) and (line_indent(block[j]) > base or not block[j].strip()):
        out.append(block[j])
        j += 1
    return out


def parse_schemas() -> dict[str, dict]:
    lines = OPENAPI.read_text().splitlines()
    start = next((i for i, line in enumerate(lines) if line.strip() == "schemas:"), -1)
    if start < 0:
        return {}
    schemas: dict[str, dict] = {}
    i = start + 1
    while i < len(lines):
        if line_indent(lines[i]) == 2 and lines[i].strip().endswith(":") and lines[i].strip() != "schemas:":
            break
        m = re.match(r"^    ([A-Za-z0-9_]+):\s*$", lines[i])
        if not m:
            i += 1
            continue
        name = m.group(1)
        j = i + 1
        while j < len(lines):
            if re.match(r"^    ([A-Za-z0-9_]+):\s*$", lines[j]) or (line_indent(lines[j]) == 2 and lines[j].strip().endswith(":")):
                break
            j += 1
        schemas[name] = parse_schema_block(lines[i:j])
        i = j
    return schemas


def parse_schema_block(block: list[str]) -> dict:
    schema = {"description": "", "required": [], "properties": []}
    in_properties = False
    current_prop = None
    for idx, line in enumerate(block):
        s = line.strip()
        ind = line_indent(line)
        if s.startswith("description: |"):
            desc = []
            j = idx + 1
            while j < len(block) and line_indent(block[j]) > ind:
                desc.append(block[j].strip())
                j += 1
            if not schema["description"]:
                schema["description"] = " ".join(desc)
        elif s.startswith("description:") and not schema["description"]:
            schema["description"] = clean_scalar(s.split(":", 1)[1])
        elif s.startswith("required:"):
            inside = s.split(":", 1)[1].strip().strip("[]")
            schema["required"] = [x.strip() for x in inside.split(",") if x.strip()]
        elif s == "properties:":
            in_properties = True
        elif in_properties and ind == 8 and re.match(r"^[A-Za-z0-9_]+:", s):
            name = s.split(":", 1)[0]
            current_prop = {"name": name, "type": "", "description": "", "required": name in schema["required"]}
            inline = s.split(":", 1)[1]
            m_type = re.search(r"type:\s*([A-Za-z0-9_-]+)", inline)
            if m_type:
                current_prop["type"] = m_type.group(1)
            m_desc = re.search(r"description:\s*([^,}]+)", inline)
            if m_desc:
                current_prop["description"] = clean_scalar(m_desc.group(1))
            m_ref = parse_schema_ref(inline)
            if m_ref:
                current_prop["type"] = m_ref
            schema["properties"].append(current_prop)
        elif current_prop and ind >= 10:
            if s.startswith("type:"):
                current_prop["type"] = clean_scalar(s.split(":", 1)[1])
            elif s.startswith("$ref:"):
                current_prop["type"] = parse_schema_ref(s) or current_prop["type"]
            elif s.startswith("format:") and current_prop["type"]:
                current_prop["type"] += f" ({clean_scalar(s.split(':', 1)[1])})"
            elif s.startswith("enum:"):
                current_prop["description"] += f" Allowed values: {clean_scalar(s.split(':', 1)[1])}."
            elif s.startswith("description: |"):
                current_prop["description"] += " See OpenAPI description."
            elif s.startswith("description:"):
                current_prop["description"] += " " + clean_scalar(s.split(":", 1)[1])
            elif s.startswith("items:") and not current_prop["type"]:
                current_prop["type"] = "array"
    return schema


def parse_error_codes() -> list[tuple[str, str, str]]:
    codes = set()
    pattern = re.compile(r'writeError\(w,\s*http\.Status([A-Za-z0-9]+),\s*"([^"]+)",\s*"([^"]*)"')
    for path in (BACKEND / "adapters" / "http").glob("handlers_*.go"):
        text = path.read_text(errors="ignore")
        for status, code, msg in pattern.findall(text):
            codes.add((status, code, msg))
    return sorted(codes, key=lambda x: (x[0], x[1]))[:180]


def sample_value(prop: dict) -> object:
    name = prop["name"]
    typ = prop.get("type", "")
    if name.endswith("_id") or name == "id":
        return f"{name}_123"
    if "email" in name:
        return "customer@example.com"
    if "phone" in name:
        return "+233501234567"
    if name == "amount" or typ.startswith("integer"):
        return 10000
    if typ.startswith("boolean"):
        return True
    if typ.startswith("array"):
        return []
    if typ.startswith("object"):
        return {}
    if "date-time" in typ or name.endswith("_at"):
        return "2026-04-24T12:00:00Z"
    if name == "currency":
        return "GHS"
    if name == "country":
        return "GH"
    if name == "method":
        return "mobile_money"
    return f"{name}_value"


def sample_body(schema_name: str, schemas: dict[str, dict]) -> str:
    if not schema_name or schema_name not in schemas:
        return ""
    schema = schemas[schema_name]
    props = schema["properties"]
    required = [p for p in props if p.get("required")]
    chosen = required or props[:5]
    if not chosen:
        chosen = props[:5]
    body = {p["name"]: sample_value(p) for p in chosen[:8]}
    return json.dumps(body, indent=2)


def tag_for(op: dict) -> str:
    if op.get("tags"):
        return op["tags"][0]
    path = op["path"]
    part = path.strip("/").split("/")
    return part[1].replace("-", " ").title() if len(part) > 1 else "General"


def auth_note(op: dict) -> str:
    names = {p["name"] for p in op["parameters"]}
    path = op["path"]
    if path in ["/healthz", "/metrics"] or path.startswith("/v1/auth/") and op["method"] in ["GET", "POST"] and any(x in path for x in ["signup", "magic-link", "verify", "recover", "oauth/login", "oauth/callback"]):
        return "Public endpoint, with route-specific rate limits where configured."
    if "XOrgId" in names or "X-Org-Id" in names:
        return "Requires organization context via X-Org-Id plus API key or session depending on route family."
    if path.startswith("/v1/platform/"):
        return "Requires dashboard session and platform admin or platform owner role."
    if path.startswith("/v1/admin/"):
        return "Requires dashboard session, organization membership, CSRF, and admin/owner role."
    if path.startswith("/v1/webhooks/"):
        return "Provider or merchant webhook route; verify provider signature where applicable."
    return "Protected according to router middleware for this route family."


def add_run(p, text: str, bold=False, color: str | None = None, size: int | None = None):
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return r


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display" if level == 1 else "Aptos"
        run.font.color.rgb = RGBColor.from_string(ACCENT if level <= 2 else GREEN)
    return p


def para(doc: Document, text: str = "", size: int = 9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if text:
        add_run(p, text, size=size)
    return p


def label_line(doc: Document, label: str, value: str, size: int = 8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    add_run(p, f"{label}: ", bold=True, color=ACCENT, size=size)
    if value:
        add_run(p, value, size=size)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=8)


def code_block(doc: Document, text: str):
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string("334155")


def configure_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(10)


def build_doc() -> None:
    go_version, deps = parse_go_mod()
    operations = parse_paths()
    schemas = parse_schemas()
    errors = parse_error_codes()
    router_text = ROUTER.read_text()

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "Reevit Backend API Documentation", bold=True, color=ACCENT, size=26)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "Detailed endpoint reference, implementation context, schemas, examples, and operations guide", color=MUTED, size=12)

    para(doc)
    label_line(doc, "Generated", date.today().isoformat())
    label_line(doc, "Backend source", "backend/adapters/http/router.go, backend/internal/docs/openapi.yaml, backend/docs/internal/*.md")
    label_line(doc, "Language/runtime", f"Go {go_version}")
    label_line(doc, "OpenAPI operations parsed", str(len(operations)))
    label_line(doc, "Router method occurrences", f"GET {router_text.count('.Get(')}, POST {router_text.count('.Post(')}, PUT {router_text.count('.Put(')}, PATCH {router_text.count('.Patch(')}, DELETE {router_text.count('.Delete(')}")
    para(doc, "This document treats router.go as the implementation source of truth and OpenAPI as the schema/operation reference. Where OpenAPI omits newer mounted routes, the implementation notes call that out.", 9)

    doc.add_page_break()
    heading(doc, "1. System Overview")
    para(doc, "Reevit is a Go backend for unified payments, billing, PSP orchestration, subscriptions, workflow automation, webhooks, fraud controls, routing, platform operations, and observability. The model is BYOK: merchants connect their own provider accounts and Reevit normalizes intent creation, confirmation, routing, webhook ingestion, retries, exports, analytics, and operational dashboards.")
    for item in [
        "HTTP uses net/http with chi routers, middleware chains, CORS origin gating, security headers, request-size limits, and provider-specific webhook payload limits.",
        "Persistence uses PostgreSQL through pgx and SQLC-generated repositories, with Goose migrations under backend/db/migrations.",
        "Redis backs idempotency, rate limiting, usage buffers, sessions/queues in selected flows, and Asynq worker queues.",
        "Authentication supports API keys, sessions, magic links, OAuth, TOTP, WebAuthn/passkeys, scoped permissions, org membership roles, and platform admin roles.",
        "Observability includes /healthz, public incidents/history, Prometheus metrics, OpenTelemetry tracing, platform logs, platform traces, failure events, connection health, and webhook delivery dashboards.",
    ]:
        bullet(doc, item)

    heading(doc, "2. Architecture and Runtime")
    for k, v in [
        ("cmd/api", "Loads config, validates production secrets, initializes database, Redis, vault, metrics, tracing, providers, services, repositories, HTTP router, and graceful server shutdown."),
        ("cmd/worker", "Runs Asynq workers/schedulers for renewals, dunning, scheduled payments, webhook fanout, and deferred side effects."),
        ("adapters/http", "Request decoding, response/error shape, chi route registration, auth/scope/rate/CSRF/plan/idempotency/mode middleware."),
        ("internal/usecase", "Business workflows for payments, connections, routing, billing, subscriptions, webhooks, exports, incidents, notifications, health, workflows, and failure reporting."),
        ("adapters/repo", "Repository implementations for ports using SQLC query files."),
        ("adapters/psp", "Provider adapters for Paystack, Hubtel, Flutterwave, Monnify, M-Pesa, Stripe, and stubs."),
        ("adapters/webhook", "Provider webhook parsers that normalize provider event shapes."),
        ("internal/infra", "Config, database, idempotency stores, HTTP clients, telemetry, logging, vault, KMS/AES secret handling."),
    ]:
        label_line(doc, k, v)

    heading(doc, "3. Libraries Used")
    for name, version in deps[:35]:
        label_line(doc, name, version, 7)

    heading(doc, "4. Global API Conventions")
    for k, v in [
        ("Base URL", "Local server is configured by SERVER_ADDR. OpenAPI lists http://localhost:8081 and https://api.reevit.io."),
        ("Versioning", "Business APIs are mounted under /v1. A few operational and legacy webhook routes are mounted at root."),
        ("Organization scope", "Protected tenant routes require X-Org-Id. Some session routes infer user context, but org-scoped work still passes through OrgScope middleware."),
        ("API key", "X-Reevit-Key: pfk_<org_id>.<secret>. Scopes such as payments:read or connections:write are enforced by route family."),
        ("Session", "Dashboard session via reevit_session cookie or Authorization: Bearer <session_token>. Session mutations use CSRF middleware."),
        ("Idempotency", "Mutating endpoints that create/update external side effects accept Idempotency-Key and replay cached responses within the configured TTL."),
        ("Money", "All amounts are integer minor units. GHS 100.00 is 10000, not 100.00."),
        ("Modes", "Live/test mode is enforced for keys, payments, routing rules, webhooks, subscriptions, and related data where repositories are mode-aware."),
        ("Errors", "Error responses follow { error: { code, message } } via writeError."),
    ]:
        label_line(doc, k, v)

    heading(doc, "5. Auth and Permission Matrix")
    scopes = [
        ("connections:read/write", "Connections, labels, audit, health, failures, credential validation, provider tests."),
        ("payments:read/write", "Payment intents, confirmation, capture, refunds, disputes, scheduled payments, payment methods, fees, analytics, exports."),
        ("fraud:read/write", "Fraud policy and audit endpoints."),
        ("webhooks:read/write", "Webhook events, replay, outbound config/test, delivery dashboard, outbound stats."),
        ("subscriptions:read/write", "Subscriptions, invoices, retry policies, dunning, exports."),
        ("api_keys:read/write", "API key lifecycle and API usage analytics."),
        ("workflows:read/write", "Integrations, templates, rules, executions, fraud rules, versioning."),
        ("routing_rules:read/write", "Routing rules, A/B tests, routing decision history."),
        ("customers:read/write", "Customer profile lifecycle and payment history."),
        ("payment_links:read/write", "Payment link lifecycle, stats, public code lookup, associated payments."),
    ]
    for k, v in scopes:
        label_line(doc, k, v)

    heading(doc, "6. Detailed Endpoint Reference")
    grouped: dict[str, list[dict]] = defaultdict(list)
    for op in operations:
        grouped[tag_for(op)].append(op)

    for tag in sorted(grouped):
        doc.add_page_break()
        heading(doc, tag, 2)
        for op in grouped[tag]:
            endpoint_title = f"{op['method']} {op['path']}"
            heading(doc, endpoint_title, 3)
            label_line(doc, "Operation ID", op["operationId"], 7)
            label_line(doc, "Purpose", op["summary"], 7)
            if op["description"]:
                label_line(doc, "Details", " ".join(op["description"].split())[:900], 7)
            label_line(doc, "Authentication", auth_note(op), 7)
            if op["parameters"]:
                label_line(doc, "Parameters", "", 7)
                for p in op["parameters"]:
                    req = "required" if str(p.get("required")) == "true" else "optional"
                    bullet(doc, f"{p.get('name')} ({p.get('in')}, {req}, {p.get('type') or 'string'}): {p.get('description') or 'No description in OpenAPI.'}")
            else:
                label_line(doc, "Parameters", "None declared in OpenAPI.", 7)
            if op["request_schema"]:
                schema = schemas.get(op["request_schema"], {})
                label_line(doc, "Request body", f"{op['request_schema']} ({'required' if op['request_required'] else 'optional'})", 7)
                props = schema.get("properties", [])[:12]
                for p in props:
                    required = "required" if p.get("required") else "optional"
                    bullet(doc, f"{p['name']} ({p.get('type') or 'any'}, {required}) {p.get('description', '').strip()}")
                example = sample_body(op["request_schema"], schemas)
                if example:
                    label_line(doc, "Example JSON", "", 7)
                    code_block(doc, example)
            else:
                label_line(doc, "Request body", "None.", 7)
            if op["responses"]:
                label_line(doc, "Responses", "", 7)
                for resp in op["responses"]:
                    extra = resp["schema"] or (f"standard {resp['ref']}" if resp["ref"] else "")
                    bullet(doc, f"{resp['code']}: {resp['description'] or 'See referenced response.'}{' -> ' + extra if extra else ''}")
            curl = make_curl(op)
            label_line(doc, "Request shape", "", 7)
            code_block(doc, curl)

    doc.add_page_break()
    heading(doc, "7. Schema Reference")
    para(doc, "This section summarizes OpenAPI component schemas used by request and response bodies. It is intentionally compact, but field names, required markers, types, and schema descriptions are preserved where available.")
    for name in sorted(schemas):
        schema = schemas[name]
        heading(doc, name, 3)
        if schema.get("description"):
            para(doc, schema["description"][:800], 8)
        props = schema.get("properties", [])
        if not props:
            para(doc, "No top-level properties declared.", 8)
            continue
        for p in props[:30]:
            required = "required" if p.get("required") else "optional"
            label_line(doc, p["name"], f"{p.get('type') or 'any'}; {required}. {p.get('description', '').strip()}", 7)

    doc.add_page_break()
    heading(doc, "8. Error Catalog")
    para(doc, "The following codes are extracted from HTTP handlers using writeError calls. Endpoint-specific OpenAPI responses appear in the endpoint reference; this catalog helps frontend and SDK clients normalize failures.")
    for status, code, msg in errors:
        label_line(doc, code, f"HTTP {status}; {msg}", 7)

    heading(doc, "9. Implementation-Only and Drift Notes")
    for item in [
        "router.go mounts additional endpoints beyond the current OpenAPI operation list, especially payment links, checkout settings, platform settings/features/incidents/logs/traces, failure events, notifications, SSE, dunning, and workflow versioning.",
        "Route availability is conditional. Stores/services passed as nil prevent selected route groups from mounting.",
        "The docs should be regenerated after router changes so SDKs, public docs, and dashboard clients do not drift from implementation.",
        "Admin and platform endpoints use session auth rather than API-key auth and may include stricter IP, 2FA, role, and audit middleware.",
        "Production must configure explicit CORS origins, secure VAULT_KEY/PASETO_KEY, provider secrets, cookie domain, webhook signing secrets, and docs basic auth if docs should be protected.",
    ]:
        bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def make_curl(op: dict) -> str:
    base = "https://api.reevit.io"
    lines = [f"curl -X {op['method']} {base}{op['path']} \\"]
    if op["path"].startswith("/v1/") and not op["path"].startswith("/v1/auth/"):
        lines.append('  -H "X-Org-Id: org_123" \\')
        lines.append('  -H "X-Reevit-Key: pfk_org_123.secret" \\')
    if op["request_schema"]:
        lines.append('  -H "Content-Type: application/json" \\')
        lines.append('  -H "Idempotency-Key: request_123" \\')
        lines.append("  -d '{...}'")
    else:
        lines[-1] = lines[-1].rstrip(" \\")
    return "\n".join(lines)


if __name__ == "__main__":
    sys.path.insert(0, str(ROOT))
    build_doc()
